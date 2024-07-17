import sys
import os
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))
from docx import Document
import re
from logs.logger import logging

class ExportOperations:
    def __init__(self, src):
        self.src = src

    def extract_product_system(self, main_table):
        try:
            product = main_table.cell(0, 1).text
            space_index = (main_table.cell(3, 1).text).find(' ')
            system = (main_table.cell(3, 1).text)[space_index + 1:]
            logging.info("Extracted product and system successfully")
            return product, system
        except Exception as e:
            logging.error(f"Error extracting product and system: {e}")
            raise

    def extract_usecase_standard(self, main_table):
        try:
            usecase = main_table.cell(1, 1).text
            standard = main_table.cell(4, 1).text
            logging.info("Extracted usecase and standard successfully")
            return usecase, standard
        except Exception as e:
            logging.error(f"Error extracting usecase and standard: {e}")
            raise

    def extract_date(self, main_table):
        try:
            text = main_table.cell(7, 1).text
            pattern = r"\b\d{2}\.\d{2}\.\d{4}\b"
            match = re.search(pattern, text)
            date = match.group(0) if match else None
            logging.info("Extracted date successfully")
            return date
        except Exception as e:
            logging.error(f"Error extracting date: {e}")
            raise

    def extract_number(self, doc):
        try:
            number = None
            for section in doc.sections:
                header = section.header
                for shape in header._element.xpath('.//w:txbxContent'):
                    for paragraph in shape.findall('.//w:p', namespaces=doc.element.nsmap):
                        number = paragraph.text
            logging.info("Extracted number successfully")
            return number
        except Exception as e:
            logging.error(f"Error extracting number: {e}")
            raise

    def extract_characteristics(self, main_table):
        try:
            characteristics = []
            cell = main_table.cell(5, 0)
            for table in cell.tables:
                table_characteristics = []
                for x, row in enumerate(table.rows):
                    if x != 0:
                        table_characteristics.append(row.cells[0].text)
                        table_characteristics.append(row.cells[1].text)
                characteristics.append(table_characteristics)
            logging.info("Extracted characteristics successfully")
            return characteristics
        except Exception as e:
            logging.error(f"Error extracting characteristics: {e}")
            raise

    def export_data(self):
        try:
            to_translate = []
            template = []

            doc = Document(self.src)
            main_table = doc.tables[0]

            product, system = self.extract_product_system(main_table)
            template.append(product)
            template.append(system)

            usecase, standard = self.extract_usecase_standard(main_table)
            to_translate.append(usecase)
            to_translate.append(standard)

            date = self.extract_date(main_table)
            template.append(date)

            number = self.extract_number(doc)
            template.append(number)

            characteristics = self.extract_characteristics(main_table)
            to_translate.extend(characteristics)

            logging.info("Exported data successfully")
            return to_translate, template
        except Exception as e:
            logging.error(f"Error exporting data: {e}")
            raise