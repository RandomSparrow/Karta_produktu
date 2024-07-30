import sys
import os
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from copy import deepcopy
from logs.logger import logging

class ImportOperations:
    def __init__(self, data):
        self.data = data

    def import_data(self, templates, save_path):
        try:
            for template, sign in templates.items():
                doc = Document(template)

                placeholders = {
                    'number': self.data.get("Number"),
                    '{{product_type}}': self.data.get("Product"),
                    '{{usecase}}': self.data.get("Usecase"),
                    '{{system}}': self.data.get("System"),
                    '{{standard}}': self.data.get("Standard"),
                    '{{date}}': self.data.get("Signature date")
                }

                doc = self._fill_placeholders(doc, placeholders)
                doc = self._fill_header(doc, placeholders)
                doc = self._fill_table(doc, self.data)

                number = (self.data.get("Number")).replace('/', '_')
                output_path = save_path + f'/{sign}_{number}_{self.data.get("Product")}_DWU_{self.data.get("Signature date")}.docx'
                doc.save(output_path)
                logging.info(f"Document saved to {output_path}")
        except Exception as e:
            logging.error(f"Error importing data: {e}")
            raise

    def _fill_placeholders(self, doc, placeholders):
        try:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for placeholder, value in placeholders.items():
                            if placeholder in cell.text:
                                cell.text = cell.text.replace(placeholder, value)
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        if placeholder == "{{date}}":
                                            run.font.name = 'Arial'
                                            run.font.size = Pt(9)
                                        if placeholder == "{{product_type}}":
                                            run.font.name = 'Arial'
                                            run.font.size = Pt(9)
                                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                            run.bold = True
                                        if placeholder in ["{{system}}", "{{standard}}"]:
                                            run.font.name = 'Arial'
                                            run.font.size = Pt(9)
                                            self._set_paragraph_format(run, "AI")
                                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                        if placeholder == "{{usecase}}":
                                            run.font.name = 'Arial'
                                            run.font.size = Pt(9)
                                            self._set_paragraph_format(run)
                                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                
            logging.info("Placeholders filled successfully")
            return doc
        except Exception as e:
            logging.error(f"Error filling placeholders: {e}")
            raise

    def _fill_header(self, doc, placeholders):
        try:
            for section in doc.sections:
                header = section.header

                for shape in header._element.xpath('.//w:txbxContent'):
                    for paragraph in shape.findall('.//w:p', namespaces=doc.element.nsmap):
                        for run in paragraph.findall('.//w:r', namespaces=doc.element.nsmap):
                            text_elem = run.find('.//w:t', namespaces=doc.element.nsmap)
                            for placeholder, value in placeholders.items():
                                if text_elem is not None and placeholder in text_elem.text:
                                    text_elem.text = text_elem.text.replace(placeholder, value)
            logging.info("Header filled successfully")
            return doc
        except Exception as e:
            logging.error(f"Error filling header: {e}")
            raise

    def _copy_table(self, table):
        # Create a deep copy of the table element
        new_tbl = deepcopy(table._element)
        return new_tbl

    def _fill_table(self, doc, data):
        try:
            main_table = doc.tables[0]
            cell_with_nested_table = main_table.cell(5, 0)

            # Get the declared performance data
            declared_performance = data.get("Declared performance")

            # Copy the existing nested table
            original_table = cell_with_nested_table.tables[0]

            # Duplicate the table for each additional set of data
            for _ in range(len(declared_performance) - 1):
                new_table_element = self._copy_table(original_table)
                cell_with_nested_table._element.append(new_table_element)
                cell_with_nested_table.add_paragraph()

            # Fill each table with the corresponding data
            for table_index, table_data in enumerate(declared_performance):
                table = cell_with_nested_table.tables[table_index]

                required_rows = len(table_data) - len(table.rows) + 1
                for _ in range(required_rows):
                  table.add_row()

                for x, element in enumerate(table_data):
                    row_cells = table.rows[x + 1].cells
                    row_cells[0].text = element.get("Essential characteristic")
                    row_cells[1].text = element.get("Performance")

                    for paragraph in row_cells[0].paragraphs:
                        for run in paragraph.runs:
                            self._set_paragraph_format(run)
                            run.font.name = 'Arial'
                            run.font.size = Pt(9)

                    for paragraph in row_cells[1].paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        for run in paragraph.runs:
                            self._set_paragraph_format(run, "AI")
                            run.font.name = 'Arial'
                            run.font.size = Pt(9)
                            

            logging.info("Table filled successfully")
            return doc
        except Exception as e:
            logging.error(f"Error filling table: {e}")
            raise

    def _set_paragraph_format(self, run, AI=None):
        if "BRAK TŁUMACZENIA" in run.text:
            run.font.color.rgb = RGBColor(255, 0, 0)  # Czerwony kolor
            run.text = run.text.replace("[AI]", "")
        elif "[AI]" in run.text:
            run.font.color.rgb = RGBColor(0, 0, 255)  # Niebieski kolor

        if AI and "[AI]" in run.text:
            run.text = run.text.replace("[AI]", "")
        